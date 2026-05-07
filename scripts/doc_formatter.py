#!/usr/bin/env python3
"""
文档格式转换工具 - 核心脚本
专门为大学生设计的作业文档格式转换工具
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
from typing import List, Dict, Tuple, Optional

class DocFormatter:
    """文档格式转换器"""
    
    def __init__(self, font_name: str = "微软雅黑", font_size: int = 11):
        """
        初始化文档格式转换器
        
        Args:
            font_name: 字体名称
            font_size: 字体大小（磅）
        """
        self.font_name = font_name
        self.font_size = Pt(font_size)
        
    def read_word_document(self, file_path: str) -> List[Dict]:
        """
        读取Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            段落列表，每个段落包含文本和样式信息
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'type': 'paragraph',
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                        'original': text
                    })
            
            # 读取表格
            for table_idx, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_cells)
                
                paragraphs.append({
                    'type': 'table',
                    'data': table_data,
                    'table_index': table_idx
                })
            
            return paragraphs
            
        except Exception as e:
            raise Exception(f"读取Word文档失败: {e}")
    
    def analyze_format_issues(self, paragraphs: List[Dict]) -> List[Dict]:
        """
        分析文档格式问题
        
        Args:
            paragraphs: 段落列表
            
        Returns:
            格式问题列表
        """
        issues = []
        
        for i, para in enumerate(paragraphs):
            if para['type'] != 'paragraph':
                continue
                
            text = para['text']
            
            # 检查标题层级
            if text.startswith('####') and '用例名称' in text:
                issues.append({
                    'type': '标题层级',
                    '位置': f"第{i+1}行",
                    '问题': '用例名称应为二级标题',
                    '建议': '将"#### 用例名称"改为"## 用例名称"'
                })
            
            # 检查列表项格式
            if text.startswith('-') and '**' in text:
                # 检查加粗格式是否正确
                if not re.match(r'^-\s*\*\*.+?\*\*\s*[:：]\s*.+', text):
                    issues.append({
                        'type': '列表格式',
                        '位置': f"第{i+1}行", 
                        '问题': '加粗列表项格式不正确',
                        '建议': '格式应为：- **标签**：内容'
                    })
            
            # 检查编号列表
            if re.match(r'^\d+\.\s*[^\s]', text):
                issues.append({
                    'type': '编号格式',
                    '位置': f"第{i+1}行",
                    '问题': '编号后缺少空格',
                    '建议': '格式应为：1. 内容'
                })
        
        return issues
    
    def format_to_markdown(self, paragraphs: List[Dict]) -> Tuple[List[str], List[Dict]]:
        """
        将文档转换为格式正确的Markdown
        
        Args:
            paragraphs: 段落列表
            
        Returns:
            (Markdown行列表, 格式问题列表)
        """
        md_lines = []
        issues = []
        
        for i, para in enumerate(paragraphs):
            if para['type'] != 'paragraph':
                continue
                
            text = para['text']
            
            # 处理主标题
            if text.startswith('### 用例描述：'):
                md_lines.append(text)
            
            # 处理分隔线
            elif text == '---':
                md_lines.append('---')
            
            # 处理用例名称
            elif text == '#### 用例名称':
                md_lines.append('## 用例名称')
            
            # 处理用例名称内容
            elif i > 0 and paragraphs[i-1]['text'] == '#### 用例名称':
                md_lines.append('### ' + text)
            
            # 处理其他子标题
            elif text in ['#### 概述', '#### 依赖', '#### 参与者', '#### 前置条件', 
                         '#### 主序列描述', '#### 可替换序列描述', 
                         '#### 非功能性需求', '#### 后置条件']:
                md_lines.append('### ' + text[4:])
            
            # 处理分支标题
            elif '分支' in text and text.startswith('####'):
                md_lines.append('#### ' + text[4:])
            
            # 处理列表项
            elif text.startswith('-'):
                # 修复加粗格式
                if '**' in text:
                    match = re.match(r'^-\s*\*\*(.+?)\*\*\s*[:：]\s*(.*)', text)
                    if match:
                        label = match.group(1)
                        content = match.group(2)
                        md_lines.append(f'- **{label}**：{content}')
                    else:
                        md_lines.append(text)
                        issues.append({
                            'type': '列表格式',
                            '位置': f"第{i+1}行",
                            '问题': '加粗列表项格式需要修复',
                            '原始文本': text
                        })
                else:
                    md_lines.append(text)
            
            # 处理编号列表
            elif re.match(r'^\d+\.', text):
                # 确保格式为：1. 内容
                formatted = re.sub(r'^(\d+)\.\s*', r'\1. ', text)
                md_lines.append(formatted)
                
                if formatted != text:
                    issues.append({
                        'type': '编号格式',
                        '位置': f"第{i+1}行",
                        '问题': '修复了编号格式',
                        '原始文本': text,
                        '修复后': formatted
                    })
            
            # 普通段落
            else:
                md_lines.append(text)
        
        return md_lines, issues
    
    def save_markdown(self, md_lines: List[str], output_path: str) -> str:
        """
        保存Markdown文件
        
        Args:
            md_lines: Markdown行列表
            output_path: 输出文件路径
            
        Returns:
            保存的文件路径
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_lines))
        
        return output_path
    
    def markdown_to_word(self, md_lines: List[str], output_path: str) -> str:
        """
        将Markdown转换为Word文档
        
        Args:
            md_lines: Markdown行列表
            output_path: 输出Word文档路径
            
        Returns:
            保存的文件路径
        """
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = self.font_size
        
        for line in md_lines:
            if not line.strip():
                doc.add_paragraph()
            
            # 处理标题
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)
            
            # 处理列表项
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                # 处理加粗文本
                parts = re.split(r'(\*\*.+?\*\*)', line[2:])
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            
            # 处理编号列表
            elif re.match(r'^\d+\.\s+', line):
                p = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s+', '', line)
                p.add_run(text)
            
            # 处理分隔线
            elif line == '---':
                p = doc.add_paragraph('_' * 50)
            
            # 普通段落
            else:
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return output_path
    
    def format_document(self, input_path: str, output_dir: str = None, 
                       formats: List[str] = None) -> Dict:
        """
        格式化文档
        
        Args:
            input_path: 输入文件路径
            output_dir: 输出目录（默认为输入文件所在目录）
            formats: 输出格式列表 ['md', 'docx']（默认两者都输出）
            
        Returns:
            处理结果信息
        """
        if formats is None:
            formats = ['md', 'docx']
        
        if output_dir is None:
            output_dir = os.path.dirname(input_path)
        
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 读取原始文档
        paragraphs = self.read_word_document(input_path)
        
        # 分析格式问题
        issues = self.analyze_format_issues(paragraphs)
        
        # 转换为Markdown
        md_lines, conversion_issues = self.format_to_markdown(paragraphs)
        issues.extend(conversion_issues)
        
        # 生成输出文件名
        input_name = os.path.splitext(os.path.basename(input_path))[0]
        results = {
            'input_file': input_path,
            'issues_found': len(issues),
            'output_files': []
        }
        
        # 保存Markdown
        if 'md' in formats:
            md_path = os.path.join(output_dir, f"格式化后的_{input_name}.md")
            self.save_markdown(md_lines, md_path)
            results['output_files'].append({
                'type': 'markdown',
                'path': md_path,
                'size': os.path.getsize(md_path)
            })
        
        # 保存Word文档
        if 'docx' in formats:
            docx_path = os.path.join(output_dir, f"作业提交_{input_name}.docx")
            self.markdown_to_word(md_lines, docx_path)
            results['output_files'].append({
                'type': 'word',
                'path': docx_path,
                'size': os.path.getsize(docx_path)
            })
        
        results['issues'] = issues
        return results
    
    def batch_format(self, input_dir: str, output_dir: str = None,
                    formats: List[str] = None) -> List[Dict]:
        """
        批量格式化文档
        
        Args:
            input_dir: 输入目录
            output_dir: 输出目录
            formats: 输出格式列表
            
        Returns:
            每个文件的处理结果列表
        """
        if output_dir is None:
            output_dir = input_dir
        
        # 查找所有Word文档
        doc_files = []
        for ext in ['.docx', '.doc']:
            doc_files.extend(list(Path(input_dir).glob(f'*{ext}')))
        
        results = []
        for doc_file in doc_files:
            try:
                result = self.format_document(str(doc_file), output_dir, formats)
                results.append(result)
            except Exception as e:
                results.append({
                    'input_file': str(doc_file),
                    'error': str(e),
                    'success': False
                })
        
        return results

def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='文档格式转换工具')
    parser.add_argument('input', help='输入文件或目录路径')
    parser.add_argument('-o', '--output', help='输出目录路径')
    parser.add_argument('-f', '--format', choices=['md', 'docx', 'both'], 
                       default='both', help='输出格式')
    parser.add_argument('-b', '--batch', action='store_true', 
                       help='批量处理目录下的所有文档')
    parser.add_argument('-r', '--report', action='store_true',
                       help='生成格式问题报告')
    
    args = parser.parse_args()
    
    # 确定输出格式
    if args.format == 'both':
        formats = ['md', 'docx']
    else:
        formats = [args.format]
    
    formatter = DocFormatter()
    
    if args.batch:
        # 批量处理
        results = formatter.batch_format(args.input, args.output, formats)
        
        print(f"批量处理完成！共处理 {len(results)} 个文档")
        print("=" * 60)
        
        for result in results:
            if 'error' in result:
                print(f"❌ {os.path.basename(result['input_file'])}: {result['error']}")
            else:
                print(f"✅ {os.path.basename(result['input_file'])}:")
                for output in result['output_files']:
                    print(f"   - {output['type']}: {output['path']}")
                
                if args.report and result['issues_found'] > 0:
                    print(f"   发现 {result['issues_found']} 个格式问题")
    
    else:
        # 单个文件处理
        result = formatter.format_document(args.input, args.output, formats)
        
        print(f"文档格式化完成！")
        print(f"输入文件: {result['input_file']}")
        print(f"发现格式问题: {result['issues_found']} 个")
        print("=" * 60)
        
        for output in result['output_files']:
            print(f"✅ {output['type'].upper()} 文件: {output['path']}")
            print(f"   文件大小: {output['size'] / 1024:.2f} KB")
        
        if args.report and result['issues_found'] > 0:
            print("\n📋 格式问题报告:")
            print("-" * 40)
            for issue in result['issues']:
                print(f"📍 {issue['位置']} - {issue['类型']}")
                print(f"   问题: {issue['问题']}")
                if '建议' in issue:
                    print(f"   建议: {issue['建议']}")
                print()

if __name__ == "__main__":
    main()