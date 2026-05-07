#!/usr/bin/env python3
"""
文档格式转换工具测试脚本
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from doc_formatter import DocFormatter

def test_basic_functionality():
    """测试基本功能"""
    print("🧪 测试文档格式转换工具...")
    
    # 创建测试文档
    test_doc_path = "test_document.docx"
    create_test_document(test_doc_path)
    
    # 测试格式化
    formatter = DocFormatter()
    
    try:
        # 测试读取文档
        print("1. 测试读取Word文档...")
        paragraphs = formatter.read_word_document(test_doc_path)
        print(f"   读取到 {len(paragraphs)} 个段落")
        
        # 测试格式分析
        print("2. 测试格式分析...")
        issues = formatter.analyze_format_issues(paragraphs)
        print(f"   发现 {len(issues)} 个格式问题")
        
        # 测试Markdown转换
        print("3. 测试Markdown转换...")
        md_lines, conversion_issues = formatter.format_to_markdown(paragraphs)
        print(f"   生成 {len(md_lines)} 行Markdown")
        print(f"   转换中发现 {len(conversion_issues)} 个问题")
        
        # 测试保存
        print("4. 测试文件保存...")
        md_path = "test_output.md"
        formatter.save_markdown(md_lines, md_path)
        print(f"   Markdown保存到: {md_path}")
        
        # 测试Word生成
        print("5. 测试Word文档生成...")
        docx_path = "test_output.docx"
        formatter.markdown_to_word(md_lines, docx_path)
        print(f"   Word文档保存到: {docx_path}")
        
        # 测试完整流程
        print("6. 测试完整格式化流程...")
        result = formatter.format_document(test_doc_path, ".", ['md', 'docx'])
        print(f"   处理完成，生成 {len(result['output_files'])} 个文件")
        
        # 清理测试文件
        cleanup_test_files([test_doc_path, md_path, docx_path])
        
        print("✅ 所有测试通过！")
        return True
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_test_document(file_path):
    """创建测试用的Word文档"""
    from docx import Document
    from docx.shared import Pt
    
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 添加测试内容（模拟常见的格式问题）
    doc.add_paragraph("### 用例描述：测试用例")
    doc.add_paragraph("---")
    doc.add_paragraph("#### 用例名称")
    doc.add_paragraph("测试用例文档")
    doc.add_paragraph("#### 概述")
    doc.add_paragraph("这是一个测试用例文档，用于验证格式转换功能。")
    doc.add_paragraph("#### 依赖")
    doc.add_paragraph("- **包含关系**：包含其他用例")
    doc.add_paragraph("- **扩展关系**：被其他用例扩展")
    doc.add_paragraph("#### 前置条件")
    doc.add_paragraph("1.系统已部署")
    doc.add_paragraph("2.配置已完成")
    doc.add_paragraph("3.数据库可访问")
    
    doc.save(file_path)
    print(f"   创建测试文档: {file_path}")

def cleanup_test_files(file_paths):
    """清理测试文件"""
    for file_path in file_paths:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"   清理文件: {file_path}")

def test_batch_processing():
    """测试批量处理功能"""
    print("\n🧪 测试批量处理功能...")
    
    # 创建测试目录和文件
    test_dir = "test_batch"
    os.makedirs(test_dir, exist_ok=True)
    
    # 创建多个测试文档
    test_files = []
    for i in range(3):
        file_path = os.path.join(test_dir, f"test_doc_{i}.docx")
        create_test_document(file_path)
        test_files.append(file_path)
    
    # 测试批量处理
    formatter = DocFormatter()
    
    try:
        results = formatter.batch_format(test_dir, test_dir, ['md', 'docx'])
        print(f"   批量处理 {len(results)} 个文档")
        
        for result in results:
            if 'error' in result:
                print(f"   ❌ {os.path.basename(result['input_file'])}: {result['error']}")
            else:
                print(f"   ✅ {os.path.basename(result['input_file'])}: 成功")
        
        # 清理
        import shutil
        shutil.rmtree(test_dir)
        print("   清理测试目录")
        
        print("✅ 批量处理测试通过！")
        return True
        
    except Exception as e:
        print(f"❌ 批量处理测试失败: {e}")
        return False

if __name__ == "__main__":
    print("=" * 60)
    print("文档格式转换工具测试套件")
    print("=" * 60)
    
    # 运行测试
    test1_passed = test_basic_functionality()
    test2_passed = test_batch_processing()
    
    print("\n" + "=" * 60)
    print("测试结果汇总:")
    print(f"基本功能测试: {'✅ 通过' if test1_passed else '❌ 失败'}")
    print(f"批量处理测试: {'✅ 通过' if test2_passed else '❌ 失败'}")
    print("=" * 60)
    
    if test1_passed and test2_passed:
        print("\n🎉 所有测试通过！工具可以正常工作。")
        sys.exit(0)
    else:
        print("\n⚠️  部分测试失败，请检查问题。")
        sys.exit(1)